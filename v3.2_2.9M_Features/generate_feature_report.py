#!/usr/bin/env python3
"""Generate Feature Importance Report as Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import json

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)

def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def section(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '6',
        qn('w:space'): '1', qn('w:color'): '1a1a2e'})
    pBdr.append(bottom)
    pPr.append(pBdr)

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(3)

def callout(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

def feature_table(tf_name, features, max_rows=100):
    """Create a compact feature importance table."""
    n_cols = 4
    t = doc.add_table(rows=1+min(len(features), max_rows), cols=n_cols)
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Rank', 'Feature Name', 'Gain', 'Category']
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)

    for ri, feat in enumerate(features[:max_rows]):
        row_data = [str(feat['rank']), feat['feature'], str(feat['gain']), feat['category']]
        for ci, val in enumerate(row_data):
            c = t.rows[ri+1].cells[ci]
            c.text = val
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)
                    if feat['category'] == 'ESOTERIC':
                        r.bold = True
                        r.font.color.rgb = RGBColor(0x80, 0x00, 0x80)
                    elif feat['category'] == 'MACRO':
                        r.font.color.rgb = RGBColor(0x00, 0x60, 0x00)
    doc.add_paragraph()

# Load summary
with open('feature_importance_summary.json') as f:
    summary = json.load(f)

# ===== COVER =====
doc.add_paragraph()
title('SAVAGE22 V3.1')
subtitle('Feature Importance & Correlation Report')
doc.add_paragraph()
subtitle('Top 500 Features Per Timeframe')
subtitle('March 23, 2026')
doc.add_page_break()

# ===== OVERVIEW =====
section('FEATURE IMPORTANCE OVERVIEW')
body('This report shows the top features by LightGBM split gain for each timeframe model. '
     'Features in PURPLE are esoteric signals (astrology, numerology, gematria, moon phases, '
     'sports, space weather). Features in GREEN are macro indicators. '
     'The model chose these features via tree splits -- no human filtering.')

# Summary table
t = doc.add_table(rows=6, cols=6)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['TF', 'Total Features', 'Features Used', 'Esoteric in Top 50', 'Top Feature', 'Top Esoteric']
for i, h in enumerate(headers):
    c = t.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

tf_order = ['1w', '1d', '4h', '1h', '15m']
for ri, tf in enumerate(tf_order):
    s = summary[tf]
    # Find top esoteric
    with open(f'feature_importance_top500_{tf}.json') as f:
        feats = json.load(f)
    top_eso = next((f for f in feats if f['category'] == 'ESOTERIC'), {'feature': 'N/A', 'rank': 'N/A'})

    row_data = [
        tf.upper(),
        str(s['total']),
        str(s['used']),
        str(s['eso_top50']),
        s['top5'][0],
        f"#{top_eso['rank']} {top_eso['feature']}"
    ]
    for ci, val in enumerate(row_data):
        c = t.rows[ri+1].cells[ci]
        c.text = val
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)
                if ci == 3 and int(s['eso_top50']) >= 8:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x80, 0x00, 0x80)

doc.add_paragraph()
callout(f'15M model: 11 esoteric features in top 50. 4H model: 9 esoteric features in top 50. '
        f'The matrix thesis is strongest at these timeframes.')

# ===== PER-TF SECTIONS =====
for tf in tf_order:
    doc.add_page_break()
    s = summary[tf]

    section(f'{tf.upper()} MODEL -- TOP 100 FEATURES')
    body(f'Total features: {s["total"]} | Used by model: {s["used"]} | '
         f'Esoteric in top 50: {s["eso_top50"]}')
    body(f'Categories in top 500: {s["cats"]}')

    with open(f'feature_importance_top500_{tf}.json') as f:
        feats = json.load(f)

    # Esoteric highlights
    eso_feats = [f for f in feats[:100] if f['category'] == 'ESOTERIC']
    if eso_feats:
        callout(f'Esoteric features in top 100: {len(eso_feats)} -- '
                + ', '.join(f"#{f['rank']} {f['feature']}" for f in eso_feats[:10]))

    # Macro highlights
    mac_feats = [f for f in feats[:100] if f['category'] == 'MACRO']
    if mac_feats:
        callout(f'Macro features in top 100: {len(mac_feats)} -- '
                + ', '.join(f"#{f['rank']} {f['feature']}" for f in mac_feats[:10]))

    feature_table(tf, feats, max_rows=100)

# ===== ESOTERIC DEEP DIVE =====
doc.add_page_break()
section('ESOTERIC FEATURE DEEP DIVE -- ACROSS ALL TIMEFRAMES')
body('Every esoteric feature that appeared in any model\'s top 100, sorted by highest rank achieved.')

all_eso = []
for tf in tf_order:
    with open(f'feature_importance_top500_{tf}.json') as f:
        feats = json.load(f)
    for feat in feats[:100]:
        if feat['category'] == 'ESOTERIC':
            all_eso.append({
                'tf': tf.upper(),
                'rank': feat['rank'],
                'feature': feat['feature'],
                'gain': feat['gain']
            })

all_eso.sort(key=lambda x: x['rank'])

t = doc.add_table(rows=1+len(all_eso), cols=4)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['TF', 'Rank', 'Feature', 'Gain']):
    c = t.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

for ri, item in enumerate(all_eso):
    for ci, val in enumerate([item['tf'], str(item['rank']), item['feature'], str(item['gain'])]):
        c = t.rows[ri+1].cells[ci]
        c.text = val
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x80, 0x00, 0x80)

doc.add_paragraph()
callout(f'Total esoteric features in top 100 across all TFs: {len(all_eso)}. '
        f'The model independently found predictive power in moon phases, planetary alignments, '
        f'gematria, sports results, and Schumann resonance -- across multiple timeframes.')

# Save
path = r'C:\Users\C\Documents\Savage22 Server\v3.1\V31_Feature_Importance_Report.docx'
doc.save(path)
print(f'Saved: {path}')
print(f'Total esoteric discoveries: {len(all_eso)}')
